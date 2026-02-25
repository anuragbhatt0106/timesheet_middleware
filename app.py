import os
import tempfile
import mimetypes
import uuid
from datetime import datetime
from flask import Flask, jsonify, request
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from docx import Document
import openpyxl
import boto3
from botocore.exceptions import ClientError, NoCredentialsError
from services.claude_service import ClaudeService

# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Configure Flask app with environment variables
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'fallback-secret-key')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# S3 Configuration - reads from environment variables
AWS_ACCESS_KEY_ID = os.environ.get('AWS_ACCESS_KEY_ID')
AWS_SECRET_ACCESS_KEY = os.environ.get('AWS_SECRET_ACCESS_KEY')
AWS_S3_BUCKET = os.environ.get('AWS_S3_BUCKET', 'saasverse-timesheet-files')
AWS_S3_REGION = os.environ.get('AWS_S3_REGION', 'ap-southeast-2')

# Initialize S3 client only if credentials are available
s3_client = None
s3_enabled = False

try:
    if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY:
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
            region_name=AWS_S3_REGION
        )
        s3_enabled = True
        print(f'S3 enabled: bucket={AWS_S3_BUCKET}, region={AWS_S3_REGION}')
    else:
        print('S3 disabled: AWS credentials not found')
except Exception as e:
    print(f'S3 init failed: {str(e)}')

# Supported file extensions and MIME types
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx', 'png', 'jpg', 'jpeg'}
ALLOWED_MIMETYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'image/png',
    'image/jpeg'
}

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_type(filename):
    """Determine file type from filename"""
    if '.' not in filename:
        return 'unknown'

    extension = filename.rsplit('.', 1)[1].lower()
    type_mapping = {
        'pdf': 'pdf',
        'docx': 'word',
        'xlsx': 'excel',
        'png': 'image',
        'jpg': 'image',
        'jpeg': 'image'
    }
    return type_mapping.get(extension, 'unknown')

def upload_to_s3(file_bytes, filename, content_type=None):
    """Upload file to S3 and return public URL. Returns None if S3 not enabled or upload fails."""
    if not s3_enabled or not s3_client:
        return None

    try:
        # Generate organized S3 key path: timesheets/YYYY/MM/uniqueid_filename
        now = datetime.utcnow()
        safe_name = secure_filename(filename)
        unique_id = uuid.uuid4().hex[:8]
        s3_key = f'timesheets/{now.year}/{now.strftime("%m")}/{unique_id}_{safe_name}'

        # Determine content type if not provided
        if not content_type:
            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            content_type_map = {
                'pdf': 'application/pdf',
                'png': 'image/png',
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'doc': 'application/msword',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }
            content_type = content_type_map.get(ext, 'application/octet-stream')

        # Upload to S3
        s3_client.put_object(
            Bucket=AWS_S3_BUCKET,
            Key=s3_key,
            Body=file_bytes,
            ContentType=content_type
        )

        # Build and return public URL
        s3_url = f'https://{AWS_S3_BUCKET}.s3.{AWS_S3_REGION}.amazonaws.com/{s3_key}'
        print(f'File uploaded to S3: {s3_url}')
        return s3_url

    except (NoCredentialsError, ClientError, Exception) as e:
        print(f'S3 upload failed: {str(e)}')
        return None

def extract_text_from_docx(file_path):
    """Extract text from Word documents"""
    try:
        doc = Document(file_path)
        text = ""

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from Word document: {str(e)}")

def extract_text_from_xlsx(file_path):
    """Extract text from Excel files"""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text = ""

        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            text += f"Sheet: {sheet_name}\n"

            for row in worksheet.iter_rows():
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value))
                if row_text:
                    text += " ".join(row_text) + "\n"

        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from Excel file: {str(e)}")

def calculate_match_status(extracted_hours, claimed_hours, confidence_score):
    """Calculate match status based on variance and confidence"""
    if confidence_score < 0.7:
        return "Low Confidence"

    variance = abs(extracted_hours - claimed_hours)

    if variance <= 0.5:
        return "Match"
    elif variance <= 2.0:
        return "Minor Mismatch"
    else:
        return "Mismatch"

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload and process with Claude AI"""
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No file provided',
                'message': 'Please provide a file in the request',
                's3_url': None,
                's3_uploaded': False
            }), 400

        file = request.files['file']

        # Check if file was actually selected
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No file selected',
                'message': 'Please select a file to upload',
                's3_url': None,
                's3_uploaded': False
            }), 400

        # Get optional parameters
        claimed_hours_str = request.form.get('claimed_hours')
        claimed_hours = None
        if claimed_hours_str:
            try:
                claimed_hours = float(claimed_hours_str)
            except ValueError:
                return jsonify({
                    'success': False,
                    'error': 'Invalid claimed_hours',
                    'message': 'claimed_hours must be a valid number',
                    's3_url': None,
                    's3_uploaded': False
                }), 400

        # Validate file type
        if not allowed_file(file.filename):
            return jsonify({
                'success': False,
                'error': 'Unsupported file type',
                'message': f'Allowed types: {", ".join(ALLOWED_EXTENSIONS)}',
                'filename': file.filename,
                's3_url': None,
                's3_uploaded': False
            }), 415

        # Secure the filename
        filename = secure_filename(file.filename)
        detected_file_type = get_file_type(filename)
        file_extension = filename.rsplit('.', 1)[1].lower()

        # Read file bytes once - reused for both AI extraction and S3 upload
        file_bytes = file.read()

        # Initialize Claude service
        claude_service = ClaudeService()

        # Create temporary file for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'_{filename}') as tmp_file:
            try:
                # Write file bytes to temp file
                tmp_file.write(file_bytes)
                tmp_file.flush()
                file_size = len(file_bytes)

                # Process file based on type
                claude_result = None

                if file_extension in ['pdf', 'png', 'jpg', 'jpeg']:
                    # Send file bytes directly to Claude
                    claude_result = claude_service.extract_timesheet_data(file_bytes, file_extension)

                elif file_extension == 'docx':
                    # Extract text from Word document and send to Claude
                    try:
                        extracted_text = extract_text_from_docx(tmp_file.name)
                        claude_result = claude_service.extract_from_text(extracted_text)
                    except Exception as e:
                        claude_result = {
                            'extracted_hours': 0,
                            'confidence_score': 0.0,
                            'summary': f'Error extracting from Word document: {str(e)}',
                            'daily_breakdown': [],
                            'anomalies': ['Word document processing failed']
                        }

                elif file_extension == 'xlsx':
                    # Extract text from Excel file and send to Claude
                    try:
                        extracted_text = extract_text_from_xlsx(tmp_file.name)
                        claude_result = claude_service.extract_from_text(extracted_text)
                    except Exception as e:
                        claude_result = {
                            'extracted_hours': 0,
                            'confidence_score': 0.0,
                            'summary': f'Error extracting from Excel file: {str(e)}',
                            'daily_breakdown': [],
                            'anomalies': ['Excel file processing failed']
                        }

                # ---- S3 UPLOAD ----
                s3_url = upload_to_s3(file_bytes, filename)

                # Build response data
                response_data = {
                    'success': True,
                    'file_name': filename,
                    'file_type': detected_file_type,
                    'file_size_bytes': file_size,
                    'extracted_hours': claude_result['extracted_hours'],
                    'confidence_score': claude_result['confidence_score'],
                    'summary': claude_result['summary'],
                    'daily_breakdown': claude_result['daily_breakdown'],
                    'anomalies': claude_result['anomalies'],
                    'approval_status': claude_result.get('approval_status'),
                    'approver_name': claude_result.get('approver_name'),
                    'resource_name': claude_result.get('resource_name'),
                    'period': claude_result.get('period'),
                    's3_url': s3_url,
                    's3_uploaded': s3_url is not None
                }

                # Add match status if claimed_hours provided
                if claimed_hours is not None:
                    match_status = calculate_match_status(
                        claude_result['extracted_hours'],
                        claimed_hours,
                        claude_result['confidence_score']
                    )
                    variance = abs(claude_result['extracted_hours'] - claimed_hours)

                    response_data.update({
                        'claimed_hours': claimed_hours,
                        'match_status': match_status,
                        'variance': variance
                    })

                return jsonify(response_data), 200

            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_file.name)
                except OSError:
                    pass  # File might already be deleted

    except Exception as e:
        return jsonify({
            'success': False,
            'error': 'Internal server error',
            'message': 'An unexpected error occurred during file processing',
            'details': str(e) if os.getenv('FLASK_ENV') == 'development' else 'Contact support',
            'extracted_hours': 0,
            'confidence_score': 0,
            'daily_breakdown': [],
            'match_status': 'Error',
            's3_url': None,
            's3_uploaded': False
        }), 500

@app.route('/api/s3-upload', methods=['POST'])
def s3_upload_only():
    """Upload a file to S3 without AI processing - useful for evidence files"""
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file_bytes = file.read()
        filename = file.filename

        s3_url = upload_to_s3(file_bytes, filename)

        if s3_url:
            return jsonify({
                'success': True,
                's3_url': s3_url,
                'filename': filename,
                'size': len(file_bytes)
            }), 200
        else:
            return jsonify({
                'success': False,
                'error': 'S3 upload failed or S3 not configured'
            }), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/')
def health():
    """Health check endpoint"""
    claude_enabled = bool(os.environ.get('ANTHROPIC_API_KEY'))
    return jsonify({
        'message': 'Timesheet Middleware is alive!',
        'status': 'healthy',
        'environment': os.getenv('FLASK_ENV', 'development'),
        'ai_enabled': claude_enabled,
        's3_enabled': s3_enabled,
        's3_bucket': AWS_S3_BUCKET if s3_enabled else None,
        'timestamp': datetime.utcnow().isoformat()
    })

@app.route('/health')
def health_check():
    """Alternative health check endpoint"""
    claude_enabled = bool(os.environ.get('ANTHROPIC_API_KEY'))
    return jsonify({
        'status': 'ok',
        'service': 'timesheet-middleware',
        'ai_enabled': claude_enabled,
        's3_enabled': s3_enabled,
        's3_bucket': AWS_S3_BUCKET if s3_enabled else None,
        'timestamp': datetime.utcnow().isoformat()
    })

@app.route('/api/status')
def api_status():
    """API status endpoint"""
    return jsonify({
        'api': 'timesheet-middleware',
        'version': '1.0.0',
        'status': 'operational',
        'supported_formats': list(ALLOWED_EXTENSIONS),
        's3_enabled': s3_enabled
    })

@app.errorhandler(413)
def too_large(error):
    """Handle file too large error"""
    return jsonify({
        'success': False,
        'error': 'File too large',
        'message': 'File size exceeds the maximum limit of 16MB',
        's3_url': None,
        's3_uploaded': False
    }), 413

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    return jsonify({
        'success': False,
        'error': 'Not Found',
        'message': 'The requested endpoint does not exist',
        'available_endpoints': ['/', '/health', '/api/status', '/api/upload (POST)', '/api/s3-upload (POST)']
    }), 404

@app.errorhandler(405)
def method_not_allowed(error):
    """Handle method not allowed errors"""
    return jsonify({
        'success': False,
        'error': 'Method Not Allowed',
        'message': 'This endpoint does not support the requested HTTP method'
    }), 405

if __name__ == '__main__':
    # This will only run when executed directly, not with gunicorn
    app.run(
        host='0.0.0.0',
        port=int(os.getenv('PORT', 5000)),
        debug=os.getenv('FLASK_ENV') == 'development'
    )
