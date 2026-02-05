import pytesseract
from PIL import Image
import cv2
import numpy as np
from pdf2image import convert_from_path
import re
import os
import tempfile
from docx import Document
import PyPDF2

class OCRService:
    def __init__(self):
        pass
    
    def extract_hours(self, file_path):
        file_extension = file_path.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension in ['png', 'jpg', 'jpeg']:
            return self._extract_from_image(file_path)
        elif file_extension in ['doc', 'docx']:
            return self._extract_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                hours = self._extract_hours_from_text(text)
                if hours is not None:
                    return hours
        except:
            pass
        
        try:
            images = convert_from_path(file_path)
            for image in images:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    image.save(tmp_file.name)
                    hours = self._extract_from_image(tmp_file.name)
                    os.unlink(tmp_file.name)
                    if hours is not None:
                        return hours
        except Exception as e:
            print(f"Error processing PDF: {e}")
        
        return None
    
    def _extract_from_image(self, file_path):
        try:
            image = cv2.imread(file_path)
            if image is None:
                return None
            
            processed_image = self._preprocess_image(image)
            
            text = pytesseract.image_to_string(processed_image)
            
            return self._extract_hours_from_text(text)
        except Exception as e:
            print(f"Error processing image: {e}")
            return None
    
    def _extract_from_word(self, file_path):
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._extract_hours_from_text(text)
        except Exception as e:
            print(f"Error processing Word document: {e}")
            return None
    
    def _preprocess_image(self, image):
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        gray = cv2.medianBlur(gray, 3)
        
        thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        
        kernel = np.ones((1,1), np.uint8)
        processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        
        return processed
    
    def _extract_hours_from_text(self, text):
        if not text:
            return None
        
        text = text.lower()
        
        hour_patterns = [
            r'total\s*hours?\s*:?\s*(\d+(?:\.\d+)?)',
            r'hours?\s*worked\s*:?\s*(\d+(?:\.\d+)?)',
            r'(\d+(?:\.\d+)?)\s*hours?',
            r'hours?\s*:?\s*(\d+(?:\.\d+)?)',
            r'total\s*:?\s*(\d+(?:\.\d+)?)\s*hrs?',
            r'(\d+(?:\.\d+)?)\s*hrs?',
            r'time\s*:?\s*(\d+(?:\.\d+)?)',
        ]
        
        for pattern in hour_patterns:
            matches = re.findall(pattern, text)
            if matches:
                try:
                    hours_value = float(matches[0])
                    if 0 <= hours_value <= 168:  # Reasonable range for weekly hours
                        return hours_value
                except ValueError:
                    continue
        
        number_pattern = r'\b(\d+(?:\.\d+)?)\b'
        numbers = re.findall(number_pattern, text)
        
        for num_str in numbers:
            try:
                num = float(num_str)
                if 1 <= num <= 168:
                    return num
            except ValueError:
                continue
        
        return None